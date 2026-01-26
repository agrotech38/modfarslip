import streamlit as st
from docx import Document
import copy
import tempfile
import os

TEMPLATE_PATH = "mod001.docx"
st.set_page_config(
    page_title="MOD JTS Label Generator",
    page_icon="📄",
    layout="centered"
)


# -----------------------------
# Replace placeholders
# -----------------------------
def replace_placeholders(doc, batch_no, counter):
    counter_text = f"[{counter}]"

    def replace_para(p):
        for run in p.runs:
            run.text = run.text.replace("{{B1}}", str(batch_no))
            run.text = run.text.replace("{{B2}}", counter_text)

    for p in doc.paragraphs:
        replace_para(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_para(p)


# -----------------------------
# Clone template page
# -----------------------------
def append_template(master, batch_no, counter):
    temp = Document(TEMPLATE_PATH)
    replace_placeholders(temp, batch_no, counter)

    for el in temp.element.body:
        master.element.body.append(copy.deepcopy(el))


# -----------------------------
# Safe filename text
# -----------------------------
def clean_filename(text):
    return text.replace("/", "-").replace("\\", "-")


# -----------------------------
# Build document
# -----------------------------
def build_file(batches):
    master = Document()
    master.element.body.clear()

    used_batches = []

    for batch in batches:
        batch_no = batch["batch_no"]
        start_counter = batch["start_counter"]
        pages = batch["pages"]

        used_batches.append(clean_filename(batch_no))

        for i in range(pages):
            counter = start_counter + (i // 2)
            append_template(master, batch_no, counter)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    master.save(tmp.name)

    # 👇 space separated (no underscore)
    filename = "MOD JTS " + " ".join(used_batches) + ".docx"

    return tmp.name, filename


# =============================
# UI
# =============================
st.title("📄 MOD JTS Label Generator")

if not os.path.exists(TEMPLATE_PATH):
    st.error("Template file 'mod001.docx' not found in repository.")
    st.stop()

num_batches = st.number_input("Number of batches", 1, 10, 1)

batches = []

for i in range(num_batches):
    st.subheader(f"Batch {i+1}")

    c1, c2, c3 = st.columns(3)

    batch_no = c1.text_input("Batch Number", key=f"b{i}")
    start_counter = c2.number_input("Jumbo Counter", value=1, key=f"c{i}")
    pages = c3.number_input("Pages", value=10, key=f"p{i}")

    batches.append({
        "batch_no": batch_no,
        "start_counter": start_counter,
        "pages": pages
    })


if st.button("Generate File"):
    file_path, filename = build_file(batches)

    with open(file_path, "rb") as f:
        st.download_button("⬇ Download DOCX", f, file_name=filename)
